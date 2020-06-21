import re

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt, Cm

from md2book.util.common import get_file_local_path
from .mdhtml import purify_remove_html

ALIGMENTS = {
	'left' : WD_ALIGN_PARAGRAPH.LEFT,
	'right' : WD_ALIGN_PARAGRAPH.RIGHT,
	'center' : WD_ALIGN_PARAGRAPH.CENTER,
	'justify' : WD_ALIGN_PARAGRAPH.JUSTIFY,
}
REGEX_IMAGES = r'\[\[IMAGE-ITEM\]\]\[\[(.*?)\]\]\[\[(.*?)\]\]'
SEP_MARGIN = Pt(20)

def create_document_styles(document):
	sep_style = document.styles.add_style('SepStyle', WD_STYLE_TYPE.PARAGRAPH)
	sep_style.base_style = document.styles['Body Text']
	# sep_style.font.size = docx.shared.Pt(14)

def post_process_docx(docx_file, target):
	docx_path = docx_file.getStrPath()
	document = docx.Document(docx_path)
	create_document_styles(document)

	alignment = target['style']['align']
	sep = target['sep']
	heading_sep = 'Heading ' + str(target['chapter-level'])
	image_alignment = 'center' if target['style']['center-blocks'] else alignment

	for i, p in enumerate(document.paragraphs):
		# Special blocks, alignment
		txt = p.text.strip()
		if txt == '[[SEP-ITEM]]':
			p.text = sep
			p.style = document.styles['SepStyle']
			p.alignment = ALIGMENTS['center']


			p.paragraph_format.space_before = SEP_MARGIN
			p.paragraph_format.space_after = SEP_MARGIN
		elif txt == '[[SKIP-ITEM]]':
			p.text = ''
			p.add_run().add_break(WD_BREAK.PAGE)
		else:
			if p.style.name == heading_sep and i > 0:
				# document.paragraphs[i-1].add_run().add_break(WD_BREAK.PAGE)
				p.paragraph_format.page_break_before = True

			if alignment in ALIGMENTS:
				p.alignment = ALIGMENTS[alignment]

		# Images
		images = []
		def detect_image(match):
			txt, url = match.group(1), match.group(2)
			images.append(url)
			return ''

		if "[[IMAGE-ITEM]]" in p.text:
			p.text = re.sub(REGEX_IMAGES, detect_image, p.text)
			r = p.add_run()
			for url in images:
				url = get_file_local_path(url)
				if url:
					img = r.add_picture(url)
			if image_alignment in ALIGMENTS:
				p.alignment = ALIGMENTS[image_alignment]

	document.save(docx_path)

# -------------------- MAKE MARDOWN-HTML CONVERTABLE -------------------- #

def purify_for_docx(code):
	REGEX_MD_IMAGE = r"!\[(.*?)\]\((.*?)\)"

	def store_html_image(match):
		txt, url = match.group(1), match.group(2)
		template = '[[IMAGE-ITEM]][[{}]][[{}]]'
		return template.format(txt, url)

	code = purify_remove_html(code)
	code = re.sub(REGEX_MD_IMAGE, store_html_image, code)
	return code